import os
import hashlib
import subprocess
from openai import OpenAI
import fitz  # PyMuPDF
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import tiktoken
from docx import Document
from bs4 import BeautifulSoup
from db import collection

client  = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
encoder = tiktoken.get_encoding("cl100k_base")

MAX_CHUNK_TOKENS = 800
MIN_CHUNK_TOKENS = 20
OVERLAP_TOKENS   = 100


def count_tokens(text: str) -> int:
    return len(encoder.encode(text))


def _make_chunk(text, file_id, page, idx, department, category, domain, filename, document_id=None, version=None):
    chunk_id = hashlib.md5(f"{file_id}_p{page}_c{idx}".encode()).hexdigest()
    return {
        "id": chunk_id,
        "text": text.strip(),
        "metadata": {
            "file_id":    file_id,
            "filename":   filename,
            "page":       page,
            "chunk_idx":  idx,
            "department": department,
            "category":   category,
            "domain":     domain, # New explicit domain scoping
            "document_id": document_id,
            "doc_version": version,
        }
    }


def _split_by_tokens(text, file_id, page, start_idx, department, category, domain, filename, document_id=None, version=None):
    tokens = encoder.encode(text)
    chunks, pos, idx = [], 0, start_idx
    while pos < len(tokens):
        end = min(pos + MAX_CHUNK_TOKENS, len(tokens))
        chunks.append(_make_chunk(encoder.decode(tokens[pos:end]),
                                  file_id, page, idx, department, category, domain, filename, document_id, version))
        pos += MAX_CHUNK_TOKENS - OVERLAP_TOKENS
        idx += 1
    return chunks


def structural_chunk(text, file_id, page, department, category, domain, filename, document_id=None, version=None):
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks, buffer, buffer_tokens, chunk_idx = [], "", 0, 0

    for para in paragraphs:
        pt = count_tokens(para)
        if pt > MAX_CHUNK_TOKENS:
            if buffer and buffer_tokens >= MIN_CHUNK_TOKENS:
                chunks.append(_make_chunk(buffer, file_id, page, chunk_idx, department, category, domain, filename, document_id, version))
                chunk_idx += 1
                buffer, buffer_tokens = "", 0
            sub = _split_by_tokens(para, file_id, page, chunk_idx, department, category, domain, filename, document_id, version)
            chunks.extend(sub); chunk_idx += len(sub)
        elif buffer_tokens + pt > MAX_CHUNK_TOKENS:
            if buffer_tokens >= MIN_CHUNK_TOKENS:
                chunks.append(_make_chunk(buffer, file_id, page, chunk_idx, department, category, domain, filename, document_id, version))
                chunk_idx += 1
            buffer, buffer_tokens = para, pt
        else:
            buffer = (buffer + "\n\n" + para).strip() if buffer else para
            buffer_tokens += pt

    if buffer and buffer_tokens >= MIN_CHUNK_TOKENS:
        chunks.append(_make_chunk(buffer, file_id, page, chunk_idx, department, category, domain, filename, document_id, version))
    return chunks


def get_embedding(text: str) -> list[float]:
    return client.embeddings.create(
        model="text-embedding-3-small", input=text[:8000]
    ).data[0].embedding


def ocr_pdf(filepath: str, file_id: str, department: str, category: str, filename: str, document_id=None, version=None) -> list:
    """Convert PDF pages to images and run OCR."""
    print(f"    [OCR] Triggered fallback for {file_id}. This may take a moment...")
    try:
        # Convert PDF to images (lowering DPI to 200 for speed, usually enough for OCR)
        images = convert_from_path(filepath, dpi=200)
        all_chunks = []
        for i, image in enumerate(images, 1):
            # Run OCR with Vietnamese and English support
            text = pytesseract.image_to_string(image, lang='vie+eng').strip()
            if len(text) < 10: continue
            
            # Since OCR loses page-level structure often, we treat each page as a single chunk 
            # or run structural chunking if text is long.
            chunks = structural_chunk(text, file_id, i, department, category, domain, filename, document_id, version)
            if not chunks and len(text) >= 10:
                chunks = [_make_chunk(text, file_id, i, 0, department, category, domain, filename, document_id, version)]
            all_chunks.extend(chunks)
        
        print(f"    [OCR] Completed: {len(all_chunks)} chunks found.")
        return all_chunks
    except Exception as e:
        print(f"    [OCR] ERROR: {e}")
        return []


def detect_metadata(filepath: str) -> tuple[str, str, str]:
    """Detects department, category, and domain from file path."""
    parts = filepath.replace("\\", "/").split("/")
    try:
        docs_idx = next(i for i, p in enumerate(parts) if p == "docs")
    except StopIteration:
        return "General", "general", "internal"
    
    rel = parts[docs_idx + 1:]
    dept = rel[0] if len(rel) > 0 else "General"
    cat = rel[1] if len(rel) > 1 else "general"
    
    # Domain detection based on keywords
    domain = "internal"
    fn = os.path.basename(filepath).lower()
    if "lao-dong" in fn or "lao_dong" in fn or "labour" in fn:
        domain = "lao_dong"
    elif "giao-thong" in fn or "giao_thong" in fn or "traffic" in fn:
        domain = "giao_thong"
    elif "doanh-nghiep" in fn or "enterprise" in fn:
        domain = "doanh_nghiep"
    elif "phap-ly" in dept.lower() or "legal" in dept.lower():
        domain = "legal"
        
    return dept, cat, domain


def ingest_pdf(filepath: str, department: str = None, category: str = None, document_id=None, version=None) -> int:
    # Use relative path as a unique ID to prevent collisions between files with same name
    docs_dir = "/data/docs"
    try:
        file_id = os.path.relpath(filepath, docs_dir).replace("\\", "/")
    except Exception:
        file_id = os.path.basename(filepath)

    filename = os.path.basename(filepath)
    if department is None or category is None:
        d, c, dom = detect_metadata(filepath)
        department = department or d
        category   = category   or c
        domain     = dom
    else:
        # If forced, we still try to detect domain
        _, _, domain = detect_metadata(filepath)

    print(f"  [Ingest] Processing: {file_id} ({department}/{category}/{domain}) v{version if version is not None else '?'}")
    
    if not os.path.exists(filepath):
        print(f"  ERROR: File not found at {filepath}")
        return 0

    all_chunks = []
    try:
        doc = fitz.open(filepath)
        print(f"    - Opened PDF. Pages: {len(doc)} | Encrypted: {doc.is_encrypted}")
        if doc.is_encrypted:
            print(f"      WARNING: This PDF is encrypted. Text extraction might fail.")
    except Exception as e:
        print(f"  ERROR opening {filepath}: {e}"); return 0

    total_chars = 0
    for page_num, page in enumerate(doc, 1):
        try:
            # Try multiple extraction methods
            text = page.get_text("text").strip()
            if not text:
                # Fallback to "blocks" if "text" is empty
                blocks = page.get_text("blocks")
                text = "\n".join([b[4] for b in blocks if b[4].strip()]).strip()
        except Exception as e:
            print(f"    [Page {page_num}] Error extracting text: {e}")
            continue
            
        total_chars += len(text)
        if len(text) < 5: 
            continue
        
        chunks = structural_chunk(text, file_id, page_num, department, category, domain, filename, document_id, version)
        if not chunks and len(text) >= 5:
            chunks = [_make_chunk(text, file_id, page_num, 0, department, category, domain, filename, document_id, version)]
            
        all_chunks.extend(chunks)

    doc.close()

    if not all_chunks:
        print(f"  WARNING: No chunks generated for {file_id} (Total chars: {total_chars})")
        if total_chars == 0:
            # Automatic OCR fallback
            all_chunks = ocr_pdf(filepath, file_id, department, category, filename, document_id, version)
            if not all_chunks:
                print(f"  HINT: OCR also failed. This file might be empty or corrupted.")
                return 0
        else:
            return 0

    print(f"    - Generated {len(all_chunks)} chunks. Starting upsert to ChromaDB...")

    BATCH = 40
    try:
        # Check if API key is set
        if not os.getenv("OPENAI_API_KEY"):
            print("  CRITICAL ERROR: OPENAI_API_KEY is not set!")
            return 0
            
        for i in range(0, len(all_chunks), BATCH):
            batch = all_chunks[i: i + BATCH]
            collection.upsert(
                ids=[c["id"] for c in batch],
                embeddings=[get_embedding(c["text"]) for c in batch],
                documents=[c["text"] for c in batch],
                metadatas=[c["metadata"] for c in batch],
            )
            print(f"    - Upserted batch {i//BATCH + 1}/{(len(all_chunks)-1)//BATCH + 1}")
    except Exception as e:
        print(f"  CRITICAL ERROR during upsert for {file_id}: {e}")
        return 0

    print(f"  → Done: {len(all_chunks)} chunks | total in DB: {collection.count()}")
    return len(all_chunks)


def extract_docx_text(filepath: str) -> str:
    doc = Document(filepath)
    parts: list[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                c = (cell.text or "").strip()
                if c:
                    cells.append(c)
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts).strip()


def extract_html_text(html: str) -> str:
    soup = BeautifulSoup(html or "", "html.parser")
    for tag in soup(["script", "style", "noscript", "header", "footer", "nav", "aside"]):
        tag.decompose()

    main = soup.find("article") or soup.find("main") or soup.body or soup
    text = main.get_text("\n", strip=True)
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    return "\n".join(lines).strip()


def ingest_text(
    text: str,
    filepath: str,
    department: str,
    category: str,
    domain: str,
    document_id=None,
    version=None,
    source_url: str | None = None,
) -> int:
    docs_dir = "/data/docs"
    try:
        file_id = os.path.relpath(filepath, docs_dir).replace("\\", "/")
    except Exception:
        file_id = os.path.basename(filepath)

    filename = os.path.basename(filepath)
    if not text or not text.strip():
        print(f"  WARNING: Empty text for {file_id}")
        return 0

    chunks = structural_chunk(
        text.strip(),
        file_id=file_id,
        page=1,
        department=department,
        category=category,
        domain=domain,
        filename=filename,
        document_id=document_id,
        version=version,
    )
    if not chunks:
        chunks = [_make_chunk(text.strip(), file_id, 1, 0, department, category, domain, filename, document_id, version)]

    # Optionally stamp source URL into metadata (for crawled docs)
    if source_url:
        for c in chunks:
            c["metadata"]["source_url"] = source_url

    print(f"    - Generated {len(chunks)} chunks. Starting upsert to ChromaDB...")

    BATCH = 40
    try:
        if not os.getenv("OPENAI_API_KEY"):
            print("  CRITICAL ERROR: OPENAI_API_KEY is not set!")
            return 0

        for i in range(0, len(chunks), BATCH):
            batch = chunks[i: i + BATCH]
            collection.upsert(
                ids=[c["id"] for c in batch],
                embeddings=[get_embedding(c["text"]) for c in batch],
                documents=[c["text"] for c in batch],
                metadatas=[c["metadata"] for c in batch],
            )
            print(f"    - Upserted batch {i//BATCH + 1}/{(len(chunks)-1)//BATCH + 1}")
    except Exception as e:
        print(f"  CRITICAL ERROR during upsert for {file_id}: {e}")
        return 0

    print(f"  → Done: {len(chunks)} chunks | total in DB: {collection.count()}")
    return len(chunks)


def ingest_docx(filepath: str, department: str = None, category: str = None, document_id=None, version=None) -> int:
    if not os.path.exists(filepath):
        print(f"  ERROR: File not found at {filepath}")
        return 0

    filename = os.path.basename(filepath)
    if department is None or category is None:
        d, c, dom = detect_metadata(filepath)
        department = department or d
        category = category or c
        domain = dom
    else:
        _, _, domain = detect_metadata(filepath)

    print(f"  [Ingest] Processing DOCX: {filename} ({department}/{category}/{domain}) v{version if version is not None else '?'}")
    try:
        text = extract_docx_text(filepath)
    except Exception as e:
        print(f"  ERROR parsing DOCX {filepath}: {e}")
        return 0

    return ingest_text(text, filepath, department, category, domain, document_id=document_id, version=version)


def ingest_doc(filepath: str, department: str = None, category: str = None, document_id=None, version=None) -> int:
    if not os.path.exists(filepath):
        print(f"  ERROR: File not found at {filepath}")
        return 0

    filename = os.path.basename(filepath)
    if department is None or category is None:
        d, c, dom = detect_metadata(filepath)
        department = department or d
        category = category or c
        domain = dom
    else:
        _, _, domain = detect_metadata(filepath)

    print(f"  [Ingest] Processing DOC: {filename} ({department}/{category}/{domain}) v{version if version is not None else '?'}")
    try:
        # Use antiword to extract text from legacy .doc files
        res = subprocess.run(["antiword", filepath], capture_output=True, text=True, check=True)
        text = res.stdout.strip()
    except Exception as e:
        print(f"  ERROR parsing DOC {filepath}: {e}")
        return 0

    return ingest_text(text, filepath, department, category, domain, document_id=document_id, version=version)


def ingest_txt(filepath: str, department: str = None, category: str = None, document_id=None, version=None) -> int:
    if not os.path.exists(filepath):
        print(f"  ERROR: File not found at {filepath}")
        return 0

    filename = os.path.basename(filepath)
    if department is None or category is None:
        d, c, dom = detect_metadata(filepath)
        department = department or d
        category = category or c
        domain = dom
    else:
        _, _, domain = detect_metadata(filepath)

    print(f"  [Ingest] Processing TXT: {filename} ({department}/{category}/{domain}) v{version if version is not None else '?'}")
    try:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
    except Exception as e:
        print(f"  ERROR reading TXT {filepath}: {e}")
        return 0

    return ingest_text(text, filepath, department, category, domain, document_id=document_id, version=version)


def ingest_html(filepath: str, department: str = None, category: str = None, document_id=None, version=None) -> int:
    if not os.path.exists(filepath):
        print(f"  ERROR: File not found at {filepath}")
        return 0

    filename = os.path.basename(filepath)
    if department is None or category is None:
        d, c, dom = detect_metadata(filepath)
        department = department or d
        category = category or c
        domain = dom
    else:
        _, _, domain = detect_metadata(filepath)

    print(f"  [Ingest] Processing HTML: {filename} ({department}/{category}/{domain}) v{version if version is not None else '?'}")
    try:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            html = f.read()
        text = extract_html_text(html)
    except Exception as e:
        print(f"  ERROR reading HTML {filepath}: {e}")
        return 0

    return ingest_text(text, filepath, department, category, domain, document_id=document_id, version=version)


def ingest_file(filepath: str, department: str = None, category: str = None, document_id=None, version=None) -> int:
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return ingest_pdf(filepath, department=department, category=category, document_id=document_id, version=version)
    if ext == ".docx":
        return ingest_docx(filepath, department=department, category=category, document_id=document_id, version=version)
    if ext == ".doc":
        return ingest_doc(filepath, department=department, category=category, document_id=document_id, version=version)
    if ext == ".txt":
        return ingest_txt(filepath, department=department, category=category, document_id=document_id, version=version)
    if ext in (".html", ".htm"):
        return ingest_html(filepath, department=department, category=category, document_id=document_id, version=version)
    print(f"  [Ingest] Skip unsupported file: {filepath}")
    return 0


def prune_orphans(docs_dir: str = "/data/docs"):
    """Delete chunks from ChromaDB whose source files no longer exist on disk."""
    results = collection.get(include=["metadatas"])
    if not results["ids"]:
        return 0

    ids_to_delete = []
    seen_files    = set()
    deleted_files = set()

    for id_, meta in zip(results["ids"], results["metadatas"]):
        file_id = meta.get("file_id")
        if not file_id: continue

        if file_id in seen_files:
            continue # already checked this file
        
        # Check if file exists on disk
        if not os.path.exists(os.path.join(docs_dir, file_id)):
            # Find all chunks for this file_id (could be optimized but safe for now)
            deleted_files.add(file_id)
        
        seen_files.add(file_id)

    if deleted_files:
        print(f"[Prune] Found {len(deleted_files)} orphan files in DB. Deleting chunks...")
        for fid in deleted_files:
            # Delete by metadata filter
            collection.delete(where={"file_id": fid})
            print(f"  - Deleted: {fid}")
    
    return len(deleted_files)


def ingest_all(docs_dir: str = "/data/docs") -> list[dict]:
    # 1. Prune orphans first
    pruned_count = prune_orphans(docs_dir)
    
    # 2. Ingest current files
    results = []
    for root, dirs, files in os.walk(docs_dir):
        if ".versions" in dirs:
            dirs.remove(".versions")
        for fname in sorted(files):
            ext = os.path.splitext(fname)[1].lower()
            if ext not in (".pdf", ".doc", ".docx", ".txt", ".html", ".htm"):
                continue
            chunks = ingest_file(os.path.join(root, fname))
            results.append({"file": fname, "chunks": chunks})
    
    total = sum(r["chunks"] for r in results)
    print(f"\nTotal: {len(results)} files, {total} chunks. (Pruned {pruned_count} files)")
    return results


if __name__ == "__main__":
    ingest_all()
